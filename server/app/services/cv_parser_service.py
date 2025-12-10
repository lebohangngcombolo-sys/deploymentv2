import os
import re
import json
from dotenv import load_dotenv
from openai import OpenAI
from cloudinary.uploader import upload as cloudinary_upload
from app.models import Requisition

# For PDF and DOCX parsing
import pdfplumber
import docx

load_dotenv()

# -----------------------------
# OpenRouter API configuration
# -----------------------------
api_key = os.getenv("OPENROUTER_API_KEY")
openai_client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=api_key,
    default_headers={"HTTP-Referer": "https://willowy-scone-c14f7c.netlify.app/"}
)

class HybridResumeAnalyzer:

    @staticmethod
    def upload_cv(file):
        """
        Upload resume file to Cloudinary and return secure URL.
        """
        try:
            result = cloudinary_upload(
                file,
                folder="candidate_cvs",
                resource_type="raw"
            )
            return result.get("secure_url")
        except Exception as e:
            print("Cloudinary Upload Error:", str(e))
            return None

    @staticmethod
    def extract_text(file):
        """
        Extract text from TXT, PDF, or DOCX file.
        """
        try:
            filename = getattr(file, "name", "") or str(file)
            ext = filename.lower().split(".")[-1]

            # TXT
            if ext in ["txt"]:
                if hasattr(file, "read"):
                    return file.read().decode("utf-8", errors="ignore")
                else:
                    with open(file, "r", encoding="utf-8", errors="ignore") as f:
                        return f.read()

            # PDF
            elif ext in ["pdf"]:
                if hasattr(file, "read"):
                    file.seek(0)
                    with pdfplumber.open(file) as pdf:
                        return "\n".join([page.extract_text() or "" for page in pdf.pages])
                else:
                    with pdfplumber.open(file) as pdf:
                        return "\n".join([page.extract_text() or "" for page in pdf.pages])

            # DOCX
            elif ext in ["docx"]:
                if hasattr(file, "read"):
                    # Save temporarily to parse
                    import tempfile
                    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    temp.write(file.read())
                    temp.close()
                    doc = docx.Document(temp.name)
                    os.unlink(temp.name)
                else:
                    doc = docx.Document(file)
                return "\n".join([p.text for p in doc.paragraphs])

            else:
                return ""
        except Exception as e:
            print("Text Extraction Error:", str(e))
            return ""

    @staticmethod
    def analyse_resume(resume_content, job_id):
        """
        Analyse resume against job description and return JSON-based results.
        """
        job = Requisition.query.get(job_id)
        if not job:
            return {
                "match_score": 0,
                "missing_skills": [],
                "suggestions": [],
                "raw_text": "Job not found"
            }

        job_description = job.description or ""

        prompt = f"""
You are an AI recruitment assistant. Compare the resume to the job description.
Return the result STRICTLY in JSON format ONLY, no extra text.
Use this exact structure:

{{
  "match_score": 0-100,
  "missing_skills": ["skill1", "skill2"],
  "suggestions": ["suggestion1", "suggestion2"]
}}

Resume:
{resume_content}

Job Description:
{job_description}
"""

        try:
            response = openai_client.chat.completions.create(
                model="openrouter/auto",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=500,
                timeout=30
            )

            raw_text = response.choices[0].message.content.strip()

            # Safe JSON parsing
            try:
                data = json.loads(raw_text)
            except json.JSONDecodeError:
                json_match = re.search(r"\{.*\}", raw_text, re.DOTALL)
                if json_match:
                    try:
                        data = json.loads(json_match.group())
                    except json.JSONDecodeError:
                        data = {}
                else:
                    data = {}

            match_score = int(data.get("match_score", 0))
            match_score = max(0, min(match_score, 100))
            missing_skills = data.get("missing_skills", [])
            suggestions = data.get("suggestions", [])

            return {
                "match_score": match_score,
                "missing_skills": missing_skills,
                "suggestions": suggestions,
                "raw_text": raw_text
            }

        except Exception as e:
            return {
                "match_score": 0,
                "missing_skills": [],
                "suggestions": [],
                "raw_text": f"Error during analysis: {str(e)}"
            }

    @staticmethod
    def upload_and_analyse(file, job_id):
        """
        Full end-to-end pipeline:
        1. Extract text from resume (TXT, PDF, DOCX)
        2. Upload to Cloudinary
        3. Analyse with AI
        Returns a dictionary with all results.
        """
        resume_content = HybridResumeAnalyzer.extract_text(file)
        if not resume_content.strip():
            return {
                "match_score": 0,
                "missing_skills": [],
                "suggestions": [],
                "raw_text": "Failed to extract text from resume"
            }

        cv_url = HybridResumeAnalyzer.upload_cv(file)
        if not cv_url:
            return {
                "match_score": 0,
                "missing_skills": [],
                "suggestions": [],
                "raw_text": "Failed to upload resume"
            }

        analysis = HybridResumeAnalyzer.analyse_resume(resume_content, job_id)
        analysis["cv_url"] = cv_url
        return analysis
